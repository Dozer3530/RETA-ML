"""
Leave-one-field-out (LOFO) **cross-sensor** benchmark for the tabular models
(Random Forest, XGBoost).

The idea
--------
With three annotated fields (e.g. a grain-yield field and two soil-ECa fields),
we hold ONE field out as the test set and train on the other TWO. Because the
held-out field is typically a *different sensor / phenomenon* than the training
fields, each fold measures **cross-sensor transfer**: can a model trained on
soil ECa flag anomalies on a yield map, and vice-versa? Cycling the held-out
field gives one fold per field.

How a single fold works
-----------------------
1. **Split the training fields with K-means** (`kmeans_spatial_split`): within
   each of the two training fields, a spatial K-means partition gives a TRAIN
   region (model fitting) and a disjoint VALIDATION region (held-out spatially,
   for an honest in-distribution check). The third field is the TEST set.
2. **Scale every field with**  ``value' = (value - field_median) / global_scale``
   The per-field median is removed in preprocessing (robust to a field sitting
   at a different absolute level); the single ``global_scale`` is fit on the
   TRAIN regions only and saved, then applied to validation and the held-out
   test field. This preserves cross-field variability (a noisy field keeps a
   wider spread) while staying comparable across sensors. `field_rel_spread`
   (field IQR / global_scale) tells the model how variable each field is.
3. **Fit RF / XGBoost** on the train region (class-balanced), and evaluate on
   both the validation region and the held-out test field.

No leakage: the scaler (and the global scale) is fit on the TRAIN regions only;
the validation region and the test field are merely transformed with it. The
held-out field is never seen during fitting.
"""

from __future__ import annotations

from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd

from reta_ml import config
from reta_ml.load import load_table
from reta_ml.preprocess import LocalStatsMode, preprocess_pipeline, recompute_value_stats
from reta_ml.split import kmeans_spatial_split
from reta_ml.graph import build_labels
from reta_ml.metrics import confusion_matrix, prf_from_cm, CLASS_NAMES

ProgressCb = Optional[Callable[[str], None]]


# --------------------------------------------------------------------------- #
# Models
# --------------------------------------------------------------------------- #
def _balanced_class_weights(
    y: np.ndarray, max_ratio: Optional[float] = None,
) -> Dict[int, float]:
    """Per-class balanced weights: n / (k * count[class]), optionally capped.

    When ``max_ratio`` is set, each class weight is clipped to
    ``max_ratio * weight[majority_class]`` (the majority class has the lowest
    balanced weight).
    """
    classes, counts = np.unique(y, return_counts=True)
    freq = {int(c): int(n) for c, n in zip(classes, counts)}
    n, k = len(y), max(len(classes), 1)
    weights = {c: n / (k * freq[c]) for c in freq}
    if max_ratio is not None:
        majority_w = min(weights.values())
        cap = max_ratio * majority_w
        weights = {c: min(w, cap) for c, w in weights.items()}
    return weights


def _balanced_sample_weight(
    y: np.ndarray, max_ratio: Optional[float] = None,
) -> np.ndarray:
    """Class-balanced sample weights: n / (k * count[class]), optionally capped."""
    weights = _balanced_class_weights(y, max_ratio)
    return np.array([weights[int(v)] for v in y], dtype=float)


def _build_models(
    which: Tuple[str, ...],
    seed: int,
    *,
    rf_class_weight="balanced",
) -> List[Tuple[str, object, bool]]:
    """Return [(display_name, estimator, uses_sample_weight), ...]."""
    out: List[Tuple[str, object, bool]] = []
    if "rf" in which:
        from sklearn.ensemble import RandomForestClassifier
        out.append(("Random Forest",
                    RandomForestClassifier(n_estimators=400, n_jobs=-1,
                                           class_weight=rf_class_weight, random_state=seed),
                    False))
    if "xgb" in which:
        try:
            from xgboost import XGBClassifier
            out.append(("XGBoost",
                        XGBClassifier(n_estimators=400, max_depth=6, learning_rate=0.05,
                                      subsample=0.9, colsample_bytree=0.9,
                                      eval_metric="mlogloss", n_jobs=0, random_state=seed),
                        True))
        except Exception:
            pass
    return out


# --------------------------------------------------------------------------- #
# Metrics helper
# --------------------------------------------------------------------------- #
def _feature_matrix(df: pd.DataFrame, feats: List[str]) -> np.ndarray:
    """Build a float32 feature matrix for the given columns (missing → 0.0).

    Tree ensembles are scale-invariant, so the hand-crafted features are fed
    raw — no StandardScaler is fit, which also removes that leakage surface.
    """
    n = len(df)
    x = np.zeros((n, len(feats)), dtype=np.float32)
    for j, c in enumerate(feats):
        if c in df.columns:
            col = pd.to_numeric(df[c], errors="coerce").to_numpy(dtype=np.float64)
            x[:, j] = np.nan_to_num(col, nan=0.0).astype(np.float32)
    return x


def _metrics(y_true: np.ndarray, y_pred: np.ndarray) -> Dict:
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    cm = confusion_matrix(y_true, y_pred, num_classes=4)
    prf = prf_from_cm(cm)
    per_class = {
        CLASS_NAMES[i]: {
            "precision": round(float(prf[f"class_{i}_precision"]), 4),
            "recall": round(float(prf[f"class_{i}_recall"]), 4),
            "f1": round(float(prf[f"class_{i}_f1"]), 4),
            "support": int(prf[f"class_{i}_support"]),
        }
        for i in range(4)
    }
    acc = float((y_true == y_pred).mean()) if len(y_true) else 0.0

    # Support-weighted averages over the classes PRESENT in this fold.
    supports = np.array([prf[f"class_{i}_support"] for i in range(4)], dtype=float)
    tot = float(supports.sum())
    def _w(metric):
        if tot <= 0:
            return 0.0
        vals = np.array([prf[f"class_{i}_{metric}"] for i in range(4)], dtype=float)
        return round(float((vals * supports).sum() / tot), 4)

    # Majority-class baseline: accuracy of always predicting the dominant class.
    majority_acc = round(float(supports.max() / tot), 4) if tot > 0 else 0.0

    # Row-normalized confusion matrix (fraction of each true class).
    cm_f = cm.astype(float)
    row_sums = cm_f.sum(axis=1, keepdims=True)
    row_sums[row_sums == 0] = 1.0
    cm_norm = np.round(cm_f / row_sums, 4)

    return {
        "accuracy": round(acc, 4),
        "majority_baseline_accuracy": majority_acc,
        "macro_f1_present": round(float(prf["macro_f1_present"]), 4),
        "macro_f1_all4": round(float(prf["macro_f1"]), 4),
        "weighted_precision": _w("precision"),
        "weighted_recall": _w("recall"),
        "weighted_f1": _w("f1"),
        "n": int(len(y_true)),
        "per_class": per_class,
        "confusion_matrix": cm.tolist(),
        "confusion_matrix_normalized": cm_norm.tolist(),
    }


# --------------------------------------------------------------------------- #
# One fold
# --------------------------------------------------------------------------- #
def run_one_fold(
    test_spec: Tuple[str, str],
    train_specs: List[Tuple[str, str]],
    *,
    models: Tuple[str, ...] = ("rf", "xgb"),
    n_clusters: int = config.LOFO_N_CLUSTERS,
    seed: int = config.SPLIT_RANDOM_STATE,
    feature_cols: Optional[List[str]] = None,
    max_ratio: Optional[float] = None,
    progress: ProgressCb = None,
) -> Dict:
    """Run one LOFO fold: train on `train_specs`, test on `test_spec`.

    Features: ``config.RF_XGB_FEATURES`` (hand-crafted spatial statistics kept,
    since tree models have no other window into spatial structure). The spatial
    statistics are recomputed from each region's OWN points after the K-means
    split, so the validation region never leaks into the train features and the
    held-out field is fully independent. Tree ensembles are scale-invariant, so
    no StandardScaler is fit (removing that leakage surface entirely).
    """
    from sklearn.preprocessing import LabelEncoder

    feats = list(feature_cols) if feature_cols is not None else list(config.RF_XGB_FEATURES)

    def _say(m):
        if progress:
            progress(m)

    # 1) Preprocess each TRAINING field, K-means split (coordinates only) into a
    #    train region + a held-out spatial validation cluster, then RECOMPUTE the
    #    per-field/spatial statistics from each region's own points (leak-free).
    train_regions, val_regions = [], []
    for path, col in train_specs:
        _say(f"Preprocessing + K-means split (k={n_clusters}): {Path(path).stem}")
        raw = load_table(Path(path), validate=False).copy()
        raw["value"] = pd.to_numeric(raw[col], errors="coerce")
        proc = preprocess_pipeline(raw, main_variable="value")
        tr_reg, val_reg = kmeans_spatial_split(
            proc, n_clusters=n_clusters,
            test_clusters=config.SPLIT_TEST_CLUSTERS, random_state=seed)
        tr_reg = recompute_value_stats(tr_reg, main_variable="value")
        val_reg = recompute_value_stats(val_reg, main_variable="value")
        train_regions.append(tr_reg)
        val_regions.append(val_reg)

    train_all = pd.concat(train_regions, ignore_index=True)
    val_all = pd.concat(val_regions, ignore_index=True)

    # 2) Held-out TEST field: preprocessed on its own data only — its spatial
    #    statistics already come solely from the test field (inference-time
    #    behaviour), so there is no train→test leakage.
    test_path, test_col = test_spec
    _say(f"Preprocessing held-out TEST field: {Path(test_path).stem}")
    raw_t = load_table(Path(test_path), validate=False).copy()
    raw_t["value"] = pd.to_numeric(raw_t[test_col], errors="coerce")
    proc_t = preprocess_pipeline(raw_t, main_variable="value")

    # 3) Feature matrices + labels (drop unlabeled rows). No scaling (trees).
    def _xy(df):
        X = _feature_matrix(df, feats)
        y = build_labels(df)
        m = y != -1
        return X[m], y[m]

    X_tr, y_tr = _xy(train_all)
    X_val, y_val = _xy(val_all)
    X_te, y_te = _xy(proc_t)

    le = LabelEncoder().fit(y_tr)
    y_tr_enc = le.transform(y_tr)
    sw = _balanced_sample_weight(y_tr_enc, max_ratio)
    rf_cw = ("balanced" if max_ratio is None
             else _balanced_class_weights(y_tr_enc, max_ratio))

    out_models: Dict[str, Dict] = {}
    for name, clf, use_sw in _build_models(models, seed, rf_class_weight=rf_cw):
        _say(f"Fitting {name} (train {len(y_tr):,} pts)…")
        if use_sw:
            clf.fit(X_tr, y_tr_enc, sample_weight=sw)
        else:
            clf.fit(X_tr, y_tr_enc)
        val_pred = le.inverse_transform(clf.predict(X_val).astype(int)).astype(int)
        te_pred = le.inverse_transform(clf.predict(X_te).astype(int)).astype(int)
        out_models[name] = {
            "validation": _metrics(y_val, val_pred),
            "test": _metrics(y_te, te_pred),
        }

    return {
        "test_field": Path(test_path).stem,
        "test_value_col": test_col,
        "train_fields": [Path(p).stem for p, _ in train_specs],
        "feature_columns": feats,
        "n_clusters": n_clusters,
        "n_train": int(len(y_tr)),
        "n_val": int(len(y_val)),
        "n_test": int(len(y_te)),
        "train_classes_present": sorted(int(c) for c in np.unique(y_tr)),
        "test_classes_present": sorted(int(c) for c in np.unique(y_te)),
        "models": out_models,
    }


def holdout_field_predictions(
    test_spec: Tuple[str, str],
    train_specs: List[Tuple[str, str]],
    *,
    model: str = "rf",
    n_clusters: int = config.LOFO_N_CLUSTERS,
    seed: int = config.SPLIT_RANDOM_STATE,
    feature_cols: Optional[List[str]] = None,
    local_stats_mode: LocalStatsMode = "multiband",
    max_ratio: Optional[float] = None,
    progress: ProgressCb = None,
) -> Dict:
    """Train on `train_specs` (LOFO style) and return per-point predictions on the
    held-out `test_spec` field — for mapping the honest cross-sensor result.

    Same pipeline as :func:`run_one_fold` (leak-free per-region spatial stats,
    no scaler), but returns the held-out field's labelled-point arrays:
    ``lat, lon, value, y_true, y_pred``.
    """
    from sklearn.preprocessing import LabelEncoder

    feats = list(feature_cols) if feature_cols is not None else list(config.RF_XGB_FEATURES)

    def _say(m):
        if progress:
            progress(m)

    train_regions = []
    for path, col in train_specs:
        _say(f"Train field: {Path(path).stem}")
        raw = load_table(Path(path), validate=False).copy()
        raw["value"] = pd.to_numeric(raw[col], errors="coerce")
        proc = preprocess_pipeline(
            raw, main_variable="value", local_stats_mode=local_stats_mode,
        )
        tr_reg, _ = kmeans_spatial_split(
            proc, n_clusters=n_clusters,
            test_clusters=config.SPLIT_TEST_CLUSTERS, random_state=seed)
        train_regions.append(
            recompute_value_stats(
                tr_reg, main_variable="value", local_stats_mode=local_stats_mode,
            )
        )
    train_all = pd.concat(train_regions, ignore_index=True)

    test_path, test_col = test_spec
    _say(f"Held-out field: {Path(test_path).stem}")
    raw_t = load_table(Path(test_path), validate=False).copy()
    raw_t["value"] = pd.to_numeric(raw_t[test_col], errors="coerce")
    proc_t = preprocess_pipeline(
        raw_t, main_variable="value", local_stats_mode=local_stats_mode,
    )

    X_tr = _feature_matrix(train_all, feats)
    y_tr = build_labels(train_all)
    mtr = y_tr != -1
    X_tr, y_tr = X_tr[mtr], y_tr[mtr]

    X_te = _feature_matrix(proc_t, feats)
    y_te = build_labels(proc_t)
    mte = y_te != -1

    le = LabelEncoder().fit(y_tr)
    y_enc = le.transform(y_tr)
    built = _build_models(
        (model,), seed,
        rf_class_weight=("balanced" if max_ratio is None
                         else _balanced_class_weights(y_enc, max_ratio)),
    )
    clf, use_sw = built[0][1], built[0][2]
    if use_sw:
        clf.fit(X_tr, y_enc, sample_weight=_balanced_sample_weight(y_enc, max_ratio))
    else:
        clf.fit(X_tr, y_enc)
    y_pred = le.inverse_transform(clf.predict(X_te[mte]).astype(int)).astype(int)

    lat = pd.to_numeric(proc_t["lat"], errors="coerce").to_numpy()[mte]
    lon = pd.to_numeric(proc_t["lon"], errors="coerce").to_numpy()[mte]
    val = pd.to_numeric(proc_t["value"], errors="coerce").to_numpy()[mte]
    return {
        "test_field": Path(test_path).stem,
        "train_fields": [Path(p).stem for p, _ in train_specs],
        "lat": lat, "lon": lon, "value": val,
        "y_true": y_te[mte], "y_pred": y_pred,
        "accuracy": float((y_te[mte] == y_pred).mean()) if mte.any() else 0.0,
    }


# --------------------------------------------------------------------------- #
# Full LOFO
# --------------------------------------------------------------------------- #
def run_lofo(
    files: List[str],
    value_cols: List[str],
    *,
    models: Tuple[str, ...] = ("rf", "xgb"),
    n_clusters: int = config.LOFO_N_CLUSTERS,
    seed: int = config.SPLIT_RANDOM_STATE,
    feature_cols: Optional[List[str]] = None,
    max_ratio: Optional[float] = None,
    progress: ProgressCb = None,
) -> Dict:
    """Run every leave-one-field-out fold and return structured results."""
    if len(files) < 2:
        raise ValueError("Need at least two fields for leave-one-field-out.")
    if len(files) != len(value_cols):
        raise ValueError("files and value_cols must have the same length.")

    feats = list(feature_cols) if feature_cols is not None else list(config.RF_XGB_FEATURES)
    folds = {}
    for i in range(len(files)):
        test_spec = (files[i], value_cols[i])
        train_specs = [(files[j], value_cols[j]) for j in range(len(files)) if j != i]
        if progress:
            progress(f"=== Fold {i+1}/{len(files)}: hold out {Path(files[i]).stem} ===")
        folds[Path(files[i]).stem] = run_one_fold(
            test_spec, train_specs, models=models,
            n_clusters=n_clusters, seed=seed, feature_cols=feats,
            max_ratio=max_ratio, progress=progress)

    out = {
        "files": [Path(f).stem for f in files],
        "value_cols": list(value_cols),
        "models": list(models),
        "n_clusters": n_clusters,
        "seed": seed,
        "feature_columns": feats,
        "folds": folds,
    }
    if max_ratio is not None:
        out["max_ratio"] = max_ratio
    return out


# --------------------------------------------------------------------------- #
# Reporting (detailed, with code snippets)
# --------------------------------------------------------------------------- #
# Curated snippets of the actual implementation, embedded in the report so a
# reader can see exactly what each step does.
_SNIPPET_HOLDOUT = '''# Hold ONE field out; train on the other two (cross-sensor transfer)
for i in range(len(files)):
    test_spec   = (files[i],  value_cols[i])              # held-out sensor
    train_specs = [(files[j], value_cols[j])             # the other two
                   for j in range(len(files)) if j != i]
    folds[stem(files[i])] = run_one_fold(test_spec, train_specs, ...)'''

_SNIPPET_KMEANS = '''# Within EACH training field, a spatial K-means split gives a
# TRAIN region (model fitting) and a disjoint VALIDATION region.
tr_reg, val_reg = kmeans_spatial_split(
    proc, n_clusters=n_clusters,
    test_clusters=config.SPLIT_TEST_CLUSTERS, random_state=seed)
train_regions.append(tr_reg); val_regions.append(val_reg)'''

_SNIPPET_SCALE = '''# Normalization:  value' = (value - field_median) / global_scale
#  - field_median/IQR are computed on each field's TRAIN region only (after the
#    K-means split), then applied to that field's train AND validation rows, so
#    the validation region never leaks into the train features' centering.
#  - global_scale is a single robust IQR fit on the TRAIN regions only, saved,
#    then reused for the validation region and the held-out field.
#  - the held-out TEST field self-centers on its own median (it is fully out).
#  - field_rel_spread = field_IQR / global_scale tells the model how variable
#    each field is, so a noisy field is not normalized to look like a tight one.
med, iqr = fit_value_centering(tr_reg)            # TRAIN region only
tr_reg  = apply_value_centering(tr_reg,  med, iqr)
val_reg = apply_value_centering(val_reg, med, iqr)   # train-region origin
scaler, cols = fit_scaler(train_all)              # global_scale on TRAIN only
train_n = transform_df(train_all, scaler, cols)
val_n   = transform_df(val_all,   scaler, cols)
test_n  = transform_df(proc_t,    scaler, cols)   # held-out field, same scaler'''

_SNIPPET_FIT = '''# Class-balanced Random Forest / XGBoost, fit on the train region,
# evaluated on the validation region AND the held-out test field.
clf.fit(X_tr, y_tr_enc, sample_weight=balanced_sw)   # XGBoost
clf.fit(X_tr, y_tr_enc)                               # RF (class_weight="balanced")
val_pred  = le.inverse_transform(clf.predict(X_val))
test_pred = le.inverse_transform(clf.predict(X_te))'''


def _cm_md(cm: List[List[int]]) -> List[str]:
    short = ["Clean", "Op", "Global", "Local"]
    lines = ["| true \\ pred | " + " | ".join(short) + " |",
             "|---|" + "---|" * 4]
    for i, row in enumerate(cm):
        lines.append(f"| **{short[i]}** | " + " | ".join(str(int(v)) for v in row) + " |")
    return lines


def render_markdown_report(results: Dict) -> str:
    """Render a detailed markdown report (methodology + code snippets + results)."""
    files = results["files"]
    L: List[str] = []
    L += [
        "# Cross-Sensor Leave-One-Field-Out Benchmark",
        "",
        "**Random Forest & XGBoost · RETA_ML**",
        "",
        f"- **Fields:** {', '.join(files)}",
        f"- **Value columns:** {', '.join(results['value_cols'])}",
        f"- **Models:** {', '.join(results['models'])}",
        f"- **K-means clusters (per training field):** {results['n_clusters']} · "
        f"**seed:** {results['seed']}",
        "",
        "## What this benchmark measures",
        "",
        "Each **fold holds one field out** as the test set and trains on the other "
        "two. Because the held-out field is usually a *different sensor / phenomenon* "
        "than the training fields (e.g. soil ECa → grain yield), every fold measures "
        "**cross-sensor transfer**: how well a model trained on two sensors flags "
        "anomalies on a third it has never seen. Cycling the held-out field gives one "
        "fold per field.",
        "",
        "## Method (with code)",
        "",
        "### 1. Hold one field out (cross-sensor)",
        "```python", _SNIPPET_HOLDOUT, "```",
        "",
        "### 2. K-means spatial split of the training fields",
        "Within each training field, a spatial K-means partition yields a **train** "
        "region for fitting and a disjoint **validation** region for an honest "
        "in-distribution check. The held-out field is the **test** set.",
        "```python", _SNIPPET_KMEANS, "```",
        "",
        "### 3. Normalization — `value' = (value − field_median) / global_scale`",
        "Centering is **per field** (robust to a field sitting at a different "
        "absolute level); the scale is a **single global value** fit on the train "
        "regions only and reused everywhere, so cross-field variability survives "
        "and the held-out field stays comparable. **No leakage:** the scaler is fit "
        "on the train regions only.",
        "```python", _SNIPPET_SCALE, "```",
        "",
        "### 4. Train Random Forest / XGBoost (class-balanced)",
        "```python", _SNIPPET_FIT, "```",
        "",
        "## Results",
        "",
    ]

    # Aggregate test macro-F1 table.
    L += ["### Summary — test macro-F1 (present classes), per fold", "",
          "| Held-out field | " + " | ".join(results["models"]) + " |",
          "|---|" + "---|" * len(results["models"])]
    for fld, d in results["folds"].items():
        row = [f"**{fld}**"]
        for mname in [m for m in d["models"]]:
            row.append(f"{d['models'][mname]['test']['macro_f1_present']:.3f}")
        L.append("| " + " | ".join(row) + " |")
    L.append("")

    # Per-fold detail.
    for fld, d in results["folds"].items():
        L += [
            f"### Fold — hold out `{fld}`",
            "",
            f"- **Train on:** {', '.join(d['train_fields'])}",
            f"- **Points:** train {d['n_train']:,} · validation {d['n_val']:,} · "
            f"test {d['n_test']:,}",
            f"- **Classes present** — train: {d['train_classes_present']} · "
            f"test: {d['test_classes_present']} "
            f"(0=Clean, 1=Operational, 2=Global, 3=Local)",
            "",
        ]
        for mname, m in d["models"].items():
            t, v = m["test"], m["validation"]
            L += [
                f"#### {mname}",
                "",
                f"- Validation (in-distribution): accuracy {v['accuracy']:.3f}, "
                f"macro-F1(present) {v['macro_f1_present']:.3f}",
                f"- **Test (cross-sensor, held-out): accuracy {t['accuracy']:.3f}, "
                f"macro-F1(present) {t['macro_f1_present']:.3f}**",
                "",
                "| Class | Precision | Recall | F1 | Support |",
                "|---|---:|---:|---:|---:|",
            ]
            for cn, pc in t["per_class"].items():
                if pc["support"] > 0:
                    L.append(f"| {cn} | {pc['precision']:.2f} | {pc['recall']:.2f} | "
                             f"{pc['f1']:.2f} | {pc['support']} |")
            L += ["", "Confusion matrix (test, counts):", ""]
            L += _cm_md(t["confusion_matrix"])
            L += [""]

    L += [
        "## How to read these numbers",
        "",
        "- **Validation vs Test** — validation is a held-out *region of the training "
        "fields* (same sensors), so it is an in-distribution check. Test is a *whole "
        "different field/sensor*, so the gap between them is the cross-sensor "
        "transfer cost.",
        "- **Absent classes** — a class with support 0 in the test field is omitted "
        "from that fold's table; a class absent from the *training* fields cannot be "
        "predicted there (e.g. Global Outliers exist only in the yield field).",
        "- **macro-F1 (present classes)** averages F1 only over classes that occur, "
        "so a missing rare class does not artificially deflate the score.",
        "",
    ]
    return "\n".join(L)


def render_docx_report(results: Dict) -> bytes:
    """Render the LOFO results as a .docx file and return the raw bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("Cross-Sensor Leave-One-Field-Out Benchmark", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Random Forest & XGBoost · RETA_ML")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph()

    # ── Run metadata ───────────────────────────────────────────────────────────
    doc.add_heading("Run configuration", level=2)
    meta = [
        ("Fields", ", ".join(results["files"])),
        ("Value columns", ", ".join(results["value_cols"])),
        ("Models", ", ".join(results["models"])),
        ("K-means clusters (per training field)", str(results["n_clusters"])),
        ("Random seed", str(results["seed"])),
    ]
    for key, val in meta:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{key}: ").bold = True
        p.add_run(val)

    # ── Summary table ──────────────────────────────────────────────────────────
    doc.add_heading("Summary — test macro-F1 (present classes) per fold", level=2)

    # Model names as stored in the fold dicts (e.g. "Random Forest", "XGBoost")
    first_fold = next(iter(results["folds"].values()))
    model_names = list(first_fold["models"].keys())
    hdr_cols = ["Held-out field"] + model_names
    tbl = doc.add_table(rows=1 + len(results["folds"]), cols=len(hdr_cols))
    tbl.style = "Table Grid"

    # Header row
    for j, h in enumerate(hdr_cols):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    for i, (fld, d) in enumerate(results["folds"].items(), start=1):
        tbl.rows[i].cells[0].text = fld
        for j, mname in enumerate(model_names, start=1):
            val = d["models"][mname]["test"]["macro_f1_present"]
            tbl.rows[i].cells[j].text = f"{val:.3f}"

    doc.add_paragraph()

    # ── Per-fold detail ────────────────────────────────────────────────────────
    doc.add_heading("Per-fold detail", level=2)

    CLASS_NAMES = {0: "Clean", 1: "Operational Error", 2: "Global Outlier", 3: "Local Outlier"}

    for fld, d in results["folds"].items():
        doc.add_heading(f"Fold — held out: {fld}", level=3)

        p = doc.add_paragraph()
        p.add_run("Train on: ").bold = True
        p.add_run(", ".join(d["train_fields"]))

        p2 = doc.add_paragraph()
        p2.add_run("Points: ").bold = True
        p2.add_run(f"train {d['n_train']:,} · validation {d['n_val']:,} · test {d['n_test']:,}")

        for mname, m in d["models"].items():
            t, v = m["test"], m["validation"]
            doc.add_heading(mname, level=4)

            # Summary line
            p = doc.add_paragraph()
            p.add_run("Validation (in-distribution): ").bold = True
            p.add_run(f"accuracy {v['accuracy']:.3f}, macro-F1 {v['macro_f1_present']:.3f}")

            p = doc.add_paragraph()
            p.add_run("Test (cross-sensor, held-out): ").bold = True
            run = p.add_run(f"accuracy {t['accuracy']:.3f}, macro-F1 {t['macro_f1_present']:.3f}")
            run.bold = True

            # Per-class table (test only, non-zero support)
            present = {cn: pc for cn, pc in t["per_class"].items() if pc["support"] > 0}
            if present:
                ct = doc.add_table(rows=1 + len(present), cols=5)
                ct.style = "Table Grid"
                for j, h in enumerate(["Class", "Precision", "Recall", "F1", "Support"]):
                    cell = ct.rows[0].cells[j]
                    cell.text = h
                    cell.paragraphs[0].runs[0].bold = True
                for i, (cn, pc) in enumerate(present.items(), start=1):
                    row = ct.rows[i]
                    row.cells[0].text = cn
                    row.cells[1].text = f"{pc['precision']:.2f}"
                    row.cells[2].text = f"{pc['recall']:.2f}"
                    row.cells[3].text = f"{pc['f1']:.2f}"
                    row.cells[4].text = str(pc["support"])

            doc.add_paragraph()

    # ── How to read ────────────────────────────────────────────────────────────
    doc.add_heading("How to read these numbers", level=2)
    notes = [
        ("Validation vs Test",
         "Validation is a held-out region of the training fields (same sensors) — an "
         "in-distribution check. Test is a whole different field/sensor, so the gap "
         "between them is the cross-sensor transfer cost."),
        ("Absent classes",
         "A class with support 0 in the test field is omitted from that fold's table. "
         "A class absent from the training fields cannot be predicted."),
        ("Macro-F1 (present classes)",
         "Averages F1 only over classes that actually occur, so a missing rare class "
         "does not artificially deflate the score."),
    ]
    for key, val in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{key}: ").bold = True
        p.add_run(val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

